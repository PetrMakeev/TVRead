#import datetime
#from gettext import find
#from operator import truediv

#from time import strptime
from docx import Document
from docx.shared import Pt
from docx.shared import Mm
import platform
import os

# инициализируем списки
lst_Ch = []
lst_Repl = []
lst_Remove = []
lst_CapsWord = []
lst_StopWord = []
lst_Programma = []
lst_txt = []
divH = 2
list_week = ['ПОНЕДЕЛЬНИК', 'ВТОРНИК', 'СРЕДА', 'ЧЕТВЕРГ', 'ПЯТНИЦА', 'СУББОТА', 'ВОСКРЕСЕНЬЕ']

vozrast = []
for i in range(0,20):
    vozrast.extend(['(' + str(i) + '+)'])
vozrast1 = []
for i in range(0,20):
    vozrast1.extend(['[' + str(i) + '+]'])
vozrast2 = []
for i in range(0,20):
    vozrast2.extend([' ' + str(i) + '+'])


god_film = []
for i in range(1900,2050):
    god_film.append(str(i))

# инициализируем документы для сохранения
doc1 = Document()
doc2 = Document()
doc3 = Document()
doc4 = Document()
doc5 = Document()
doc6 = Document()
doc7 = Document()
docN = Document()

# возвращаем название канала по имени файла
def id_to_name_Ch(name_Ch_id):
    for el in lst_Ch:
        if el[0]==name_Ch_id:
            Rezult = el[1]
            break
        else:
            Rezult = name_Ch_id
    return Rezult

# возращаем надо делать корректировку времени или нет
def fl_to_name_Ch(name_Ch_id):
    Rezult = False
    for el in lst_Ch:
        if el[0]==name_Ch_id:
            Rezult = (el[2]=='1')
            break
    return Rezult

# отрисовка прогресс спИна
def progressSpin(i):
    tmp_i = i % 4
    if tmp_i == 0:
        tmp_s = '-'
    if tmp_i == 1:
        tmp_s = '\\'
    if tmp_i == 2:
        tmp_s =' |'
    if tmp_i == 3:
        tmp_s = '/'
    return tmp_s

#подготовка экрана
def set_scr():
    os.system('cls' if os.name == 'nt' else 'clear')
    print('Парсер для обработки программ телеканалов ')
    print('Макеев Петр тел. 8-912-34-161-34 ')
    print('-----------------------------------------------')

# возвращаем строку от первой кавычки до точки
def str_kav_dot(str_prog):
    # определяем в строке позиции кавычки ёлочки
    pos1 = str_prog.find('«')
    pos2 = str_prog.find('»')
    Rezult = ''

    if (pos1 > -1) and (pos2 > pos1+2) :
        # если кавычки найдены ищем точку 
        pos3 = str_prog.find('.', pos1, pos2)
        if pos3 > -1:
            # точка найдена - вырезаем строку
            Rezult = str_prog[pos1:pos3]

    return Rezult



# поправка на часовой пояс
def timeDiv(strTime, fl_change):
    if fl_change:
        if strTime.find(':')>-1 or  strTime.find('.')>-1:
            if ':' in strTime:
                strHourN = int(strTime.split(':')[0]) + divH
            else:
                strHourN = int(strTime.split('.')[0]) + divH
            if strHourN>23 : strHourN = strHourN - 24
            if strHourN<0  : strHourN = 24 - strHourN
            if strHourN<10:
                Rezult = '0' + str(strHourN) + '.' + strTime[-2:]
            else:
                Rezult = str(strHourN) + '.' + strTime[-2:]
        else:
            Rezult = strTime
    else:
        if strTime.find(':')>-1 or  strTime.find('.')>-1:
            if ':' in strTime:
                strHourN = int(strTime.split(':')[0])
            else:
                strHourN = int(strTime.split('.')[0])
            if strHourN<10:
                Rezult = '0' + str(strHourN) + '.' + strTime[-2:]
            else:
                Rezult = str(strHourN) + '.' + strTime[-2:]
        else:
            Rezult = strTime
    return Rezult


# готовим список каналов и справочники обработок
def txt_to_list(path_prog, path_prog_in):

    global lst_txt
    global divH

    # создаем список файлов источников
    list_in = os.listdir(path_prog_in)
    for name_files_in in list_in:
        if name_files_in[-3:]=='txt':
            lst_txt.append(path_prog_in + name_files_in)

    # Справочник с телеканалами Chanenl.txt
    try:
        with open(path_prog + 'Channel.txt', 'r', encoding='Windows-1251') as file_r:
            str_txt_ch = file_r.readlines()
    except:
        # cправочник каналов Channel.txt недоступен
        print('Не найден файл со списком каналов - Channel.txt!')
        exit()

    # заполняем справочник каналов 
    for i, el in enumerate(str_txt_ch):
        if not (el[0] == '#' or el.strip()==''):
            lst_Ch.append(el.replace('\n','').split('|'))
            print('Считываем настройки - ' + progressSpin(i), end='\r')


    # Справочник с программами Programma.txt
    try:
        with open('Programma.txt', 'r', encoding='Windows-1251') as file_r:
            str_programma = file_r.readlines()
    except:
        # cправочник каналов Programma.txt недоступен
        print('Не найден файл с программами - Programma.txt!')
        exit()

    # сбрасываем текущий канал
    currCH = ''
    # заполняем справочник программ 
    for i, el in enumerate(str_programma):
        if not (el[0] == '#' or el.replace('\n','').strip()==''):
            if el.split('|')[0] == 'CH':
                currCH = el.replace('\n','').split('|')[1].strip()
                continue
            rep0 = el.split('|')[0].strip()
            rep1 = el.replace('\n','').split('|')[1].strip()
            rep2 = currCH
            lst_Programma.append([rep0, rep1, rep2])
            print('Считываем настройки - ' + progressSpin(i), end='\r')



    # Справочник с заменами Replace.txt
    try:
        with open('Replace.txt', 'r', encoding='Windows-1251') as file_r:
            str_txt_rpl = file_r.readlines()
    except:
        # cправочник каналов Replace.txt недоступен
        print('Не найден файл со списком замен - Replace.txt!')
        exit()

    # сбрасываем текущий канал
    currCH = ''
    # заполняем справочник замен 
    for i, el in enumerate(str_txt_rpl):
        if not (el[0] == '#' or el.replace('\n','').strip()==''):
            if el.split('|')[0] == 'CH':
                currCH = el.replace('\n','').split('|')[1].strip()
                continue
            rep0 = el.split('|')[0].strip()
            rep1 = el.split('|')[1].strip()
            rep2 = currCH
            lst_Repl.append([rep0, rep1, rep2])
            print('Считываем настройки - ' + progressSpin(i), end='\r')


    # справочник удалений с заменами Remove.txt
    try:
        with open('Remove.txt', 'r', encoding='Windows-1251') as file_r:
            str_txt_rem = file_r.readlines()
    except:
        # cправочник каналов Remove.txt недоступен
        print('Не найден файл со списком удалений - Remove.txt!')
        exit()

    for i, el in enumerate(str_txt_rem):
        if not (el[0] == '#' or el.strip()==''):
            lst_Remove.append(el.replace('\n','').strip())
            print('Считываем настройки - ' + progressSpin(i), end='\r')


    # справочник с исключениями CapsWord.txt
    try:
        with open('CapsWord.txt', 'r', encoding='Windows-1251') as file_r:
            str_txt_cpsl = file_r.readlines()
    except:
        # cправочник каналов CapsWord.txt недоступен
        print('Не найден файл со списком удалений - CapsWord.txt!')
        exit()

    # заполняем справочник капслоков 
    for i, el in enumerate(str_txt_cpsl):
        if not (el[0] == '#' or el.strip()==''):
            lst_CapsWord.append(el.replace('\n','').strip())
            print('Считываем настройки - ' + progressSpin(i), end='\r')


    # Справочник с исключениями StopWord.txt
    try:
        with open('StopWord.txt', 'r', encoding='Windows-1251') as file_r:
            str_txt_stopw = file_r.readlines()
    except:
        # cправочник каналов StopWord.txt недоступен
        print('Не найден файл со списком удалений - StopWord.txt!')
        exit()

    # заполняем справочник стоп слов 
    for i, el in enumerate(str_txt_stopw):
        if not (el[0] == '#' or el.strip()==''):
            lst_StopWord.append(el.replace('\n','').strip())
            print('Считываем настройки - ' + progressSpin(i), end='\r')

    print('Считываем настройки - ВЫПОЛНЕНО!')



# готовим  списки телепрограмм по дням недели для наполнения
def fill_Day():

    global lst_D1
    global lst_D2
    global lst_D3
    global lst_D4
    global lst_D5
    global lst_D6
    global lst_D7

    lst_D1 = [['~']]
    for el in lst_Ch:
        lst_D1.append([el[1]])
    lst_D1.remove(['~'])
    lst_D2 = [['~']]
    for el in lst_Ch:
        lst_D2.append([el[1]])
    lst_D2.remove(['~'])
    lst_D3 = [['~']]
    for el in lst_Ch:
        lst_D3.append([el[1]])
    lst_D3.remove(['~'])
    lst_D4 = [['~']]
    for el in lst_Ch:
        lst_D4.append([el[1]])
    lst_D4.remove(['~'])
    lst_D5 = [['~']]
    for el in lst_Ch:
        lst_D5.append([el[1]])
    lst_D5.remove(['~'])
    lst_D6 = [['~']]
    for el in lst_Ch:
        lst_D6.append([el[1]])
    lst_D6.remove(['~'])
    lst_D7 = [['~']]
    for el in lst_Ch:
        lst_D7.append([el[1]])
    lst_D7.remove(['~'])



# сканируем файлы txt и формируем  программы
def txt_to_prog(path_prog):

    progressInt = 0
    for name_files in lst_txt:
        # открываем файлы txt построчно сохраняем в список
        with open(name_files, 'r', encoding='Windows-1251') as file_r:
            str_txt = file_r.readlines()

        # устанавливаем первый день недели, канал, программу
        name_Day = 'None'
        name_Ch_id = os.path.basename(name_files).split('.')[0]
        name_Ch = id_to_name_Ch(name_Ch_id)

        name_Pr = ''

        # обрабатываем список с программой из файла источника 
        fl_merge = False
        fl_merge2 = False
        for str_line in str_txt:
            # проверяем первый знак 
            str_tmp = str_line[:-1].strip()
            # и пустые строки
            if len(str_tmp)<3: continue

            print('Загружаем телепрограммы - ' + progressSpin(progressInt), end='\r')
            progressInt += 1

            str_sub_lineD = str_tmp.split(',',2)
            if str_sub_lineD[0].upper() in list_week:
                # в строке день недели
                name_Day = str_sub_lineD[0].upper()
            else:
                # в строке программа
                # отделяем время программы от названия программы
                if len(str_sub_lineD)>0:
                    str_sub_lineD = str_tmp.split(' ',1)

                    # если начинается строка с буквы или кавычки
                    if str_sub_lineD[0][0].isalpha() or str_sub_lineD[0][0] == '"'  or str_sub_lineD[0][0] == '«'  or name_Pr=='None' :
                        if fl_merge:
                            fl_merge2 = True
                        fl_merge = True
                    else:
                        fl_merge = False
                        fl_merge2 = False

                    # если второе объединение строки то соединеям |  а не   пробелом
                    if fl_merge:
                        if fl_merge2:
                            name_Pr = name_Pr + ' ' + str_tmp
                        else:
                            name_Pr = name_Pr + '|' + str_tmp
                    else:
                        if len(str_sub_lineD)==1:
                            name_Pr = timeDiv(str_sub_lineD[0], fl_to_name_Ch(name_Ch_id))
                        else:
                            name_Pr = timeDiv(str_sub_lineD[0], fl_to_name_Ch(name_Ch_id)) + '|' + str_sub_lineD[1]

                    # собираем список [канал, [программа]]
                    if name_Day == 'ПОНЕДЕЛЬНИК':
                        for i, el in enumerate(lst_Ch):
                            if el[1]==name_Ch:
                                if fl_merge:
                                    if fl_merge2:
                                        lst_D1[i][-1] = lst_D1[i][-1] + ' ' + str_tmp
                                    else:
                                        lst_D1[i][-1] = name_Pr
                                else:
                                    lst_D1[i].append(name_Pr.strip())

                    elif name_Day == 'ВТОРНИК':
                        for i, el in enumerate(lst_Ch):
                            if el[1]==name_Ch:
                                if fl_merge:
                                    lst_D2[i][-1] = lst_D2[i][-1] + '|' + str_tmp
                                else:
                                    lst_D2[i].append(name_Pr.strip())

                    elif name_Day == 'СРЕДА':
                        for i, el in enumerate(lst_Ch):
                            if el[1]==name_Ch:
                                if fl_merge:
                                    lst_D3[i][-1] = lst_D3[i][-1] + '|' + str_tmp
                                else:
                                    lst_D3[i].append(name_Pr.strip())

                    elif name_Day == 'ЧЕТВЕРГ':
                        for i, el in enumerate(lst_Ch):
                            if el[1]==name_Ch:
                                if fl_merge:
                                    lst_D4[i][-1] = lst_D4[i][-1] + '|' + str_tmp
                                else:
                                    lst_D4[i].append(name_Pr.strip())

                    elif name_Day == 'ПЯТНИЦА':
                         for i, el in enumerate(lst_Ch):
                            if el[1]==name_Ch:
                                if fl_merge:
                                    lst_D5[i][-1] = lst_D5[i][-1] + '|' + str_tmp
                                else:
                                    lst_D5[i].append(name_Pr.strip())

                    elif name_Day == 'СУББОТА':
                        for i, el in enumerate(lst_Ch):
                            if el[1]==name_Ch:
                                if fl_merge:
                                    lst_D6[i][-1] = lst_D6[i][-1] + '|' + str_tmp
                                else:
                                    lst_D6[i].append(name_Pr.strip())

                    elif name_Day == 'ВОСКРЕСЕНЬЕ':
                        for i, el in enumerate(lst_Ch):
                            if el[1]==name_Ch:
                                if fl_merge:
                                    lst_D7[i][-1] = lst_D7[i][-1] + '|' + str_tmp
                                else:
                                    lst_D7[i].append(name_Pr.strip())
    print('Загружаем телепрограммы - ВЫПОЛНЕНО!')


def analiz_prog_day(lst_day):

    progressInt = 0
    fl_Caps = False

    # перебираем программы 
    for l, el_D in enumerate(lst_day):
        # определяем нужна обработка Капса или нет
        for ll, el_ch in enumerate(lst_Ch):
            if el_D[0] == el_ch[1]:
                fl_Caps = (el_ch[3]) == '1'
                break
        for i, el_Pr in enumerate(el_D):
            if i>0:
                # сохранение телепрограмм 
               lst_day[l][i] = el_Pr.split('|', 1)[0] + '|' + analiz_prog_day_in(el_Pr.split('|', 1)[1], el_D[0], fl_Caps)

    print('Анализ телепрограмм - ' + progressSpin(progressInt) , end='\r')
    progressInt +=  1

    return lst_day

# перебираем списки программ для вызова анализа
def analiz_prog():

    progressInt = 0
    fl_Caps = False

    global lst_D1
    global lst_D2
    global lst_D3
    global lst_D4
    global lst_D5
    global lst_D6
    global lst_D7

    lst_D1 = analiz_prog_day(lst_D1)
    print('Анализ телепрограмм - ' + progressSpin(progressInt) , end='\r')
    progressInt +=  1

    lst_D2 = analiz_prog_day(lst_D2)
    print('Анализ телепрограмм - ' + progressSpin(progressInt) , end='\r')
    progressInt +=  1

    lst_D3 = analiz_prog_day(lst_D3)
    print('Анализ телепрограмм - ' + progressSpin(progressInt) , end='\r')
    progressInt +=  1

    lst_D4 = analiz_prog_day(lst_D4)
    print('Анализ телепрограмм - ' + progressSpin(progressInt) , end='\r')
    progressInt +=  1

    lst_D5 = analiz_prog_day(lst_D5)
    print('Анализ телепрограмм - ' + progressSpin(progressInt) , end='\r')
    progressInt +=  1

    lst_D6 = analiz_prog_day(lst_D6)
    print('Анализ телепрограмм - ' + progressSpin(progressInt) , end='\r')
    progressInt +=  1

    lst_D7 = analiz_prog_day(lst_D7)

    print('Анализ телепрограмм - ВЫПОЛНЕНО')


# сохранение сведений о телепрограмме 
def save_prog(doc_, doc_N, el_Pr_, el_D_, i, str_prog_, str_prog_N):
    if i<1:
        if len(el_D_) > 1:
            str_prog_= str_prog_ + 'STYLE K ' + el_Pr_.upper() + '\n'
            str_prog_N=  str_prog_N + ' ' + el_Pr_.upper() + '\n'
            # doc
            paragraph = doc_.add_paragraph('STYLE K ' + el_Pr_.upper() )
            paragraph.paragraph_format.space_before = Mm(0)
            paragraph.paragraph_format.space_after = Mm(0)

            paragraph = doc_N.add_paragraph('-------------------------------')
            paragraph.paragraph_format.space_before = Mm(0)
            paragraph.paragraph_format.space_after = Mm(0)
            paragraph = doc_N.add_paragraph(' ' + el_Pr_.upper() )
            paragraph.paragraph_format.space_before = Mm(0)
            paragraph.paragraph_format.space_after = Mm(0)
        else:
            str_prog_ = str_prog_+ 'STYLE K ' + el_Pr_.upper() + ' !!!! ПРОПУЩЕН !!!!' +  '\n'
            str_prog_N =  str_prog_N + ' ' + el_Pr_.upper() + ' !!!! ПРОПУЩЕН !!!!' +  '\n'
            # doc
            paragraph = doc_.add_paragraph('STYLE K ' + el_Pr_.upper() + ' !!!! ПРОПУЩЕН !!!!')
            paragraph.paragraph_format.space_before = Mm(0)
            paragraph.paragraph_format.space_after = Mm(0)

            paragraph = doc_N.add_paragraph('-------------------------------')
            paragraph.paragraph_format.space_before = Mm(0)
            paragraph.paragraph_format.space_after = Mm(0)
            paragraph = doc_N.add_paragraph(' ' + el_Pr_.upper() + ' !!!! ПРОПУЩЕН !!!!')
            paragraph.paragraph_format.space_before = Mm(0)
            paragraph.paragraph_format.space_after = Mm(0)

    else:

        #анализируем и обрабатываем строку
        repl_str_prog = el_Pr_.split('|',1)[1]

        if '^' in repl_str_prog:
            vozrast_id = repl_str_prog.split('^')[-1]
        else:
            vozrast_id =''

        str_prog_ = str_prog_ + el_Pr_.split('|',1)[0] + ' ' + repl_str_prog + '\n'
        str_prog_N = str_prog_N + el_Pr_.split('|',1)[0] + ' ' + repl_str_prog + '\n'

        # doc
        paragraph = doc_.add_paragraph()
        paragraph.add_run(el_Pr_.split('|',1)[0] ).bold = True
        paragraph.add_run(' ' + repl_str_prog.split('^')[0]).bold = False
        if not vozrast_id == '':
            paragraph.add_run(vozrast_id).font.superscript = True
        paragraph.paragraph_format.space_before = Mm(0)
        paragraph.paragraph_format.space_after = Mm(0)

        paragraph = doc_N.add_paragraph()
        paragraph.add_run(el_Pr_.split('|',1)[0] ).bold = True
        paragraph.add_run(' ' + repl_str_prog.split('^')[0]).bold = False
        if not vozrast_id == '':
            paragraph.add_run(vozrast_id).font.superscript = True
        paragraph.paragraph_format.space_before = Mm(0)
        paragraph.paragraph_format.space_after = Mm(0)
    return [str_prog_, str_prog_N]



# экспорт в файлы
def exp_prog(path_prog_out):


    # задаем стиль текста по умолчанию
    style = doc1.styles['Normal']
    # название шрифта
    style.font.name = 'Arial'
    # размер шрифта
    style.font.size = Pt(12)

    progressInt = 0

    # понедельник
    str_prog1 = 'STYLE D \n'
    str_progN = '-------------------------------\n' + ' ПОНЕДЕЛЬНИК\n'

    # doc
    paragraph = doc1.add_paragraph('STYLE D ')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)

    paragraph = docN.add_paragraph('-------------------------------')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)
    paragraph = docN.add_paragraph(' ПОНЕДЕЛЬНИК')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)

    for el_D in lst_D1:
        # перебираем программы
        for i, el_Pr in enumerate(el_D):

            # сохранение телепрограмм
            lst_tmp = save_prog(doc1, docN, el_Pr, el_D, i, str_prog1, str_progN)
            str_prog1 = lst_tmp[0]
            str_progN = lst_tmp[1]

            print(('Сохранение телепрограмм Понедельника - ' + progressSpin(progressInt)).ljust(60, " ") , end='\r')
            progressInt +=  1


     # вторник
    str_prog2 = 'STYLE D \n'
    str_progN =  str_progN + '-------------------------------\n' + ' ВТОРНИК\n'

    # doc
    paragraph = docN.add_paragraph('-------------------------------')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)
    paragraph = docN.add_paragraph(' ВТОРНИК')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)

    paragraph = doc2.add_paragraph('STYLE D ')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)

    for el_D in lst_D2:
        # перебираем программы
        for i, el_Pr in enumerate(el_D):

            # сохранение телепрограмм
            lst_tmp = save_prog(doc2, docN, el_Pr, el_D, i, str_prog2, str_progN)
            str_prog2 = lst_tmp[0]
            str_progN = lst_tmp[1]

            print(('Сохранение телепрограмм Вторника - ' + progressSpin(progressInt)).ljust(60, " ") , end='\r')
            progressInt +=  1


    # вторник
    str_prog3 = 'STYLE D \n'
    str_progN = str_progN + '-------------------------------\n' + ' СРЕДА\n'


    # doc
    paragraph = docN.add_paragraph('-------------------------------')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)
    paragraph = docN.add_paragraph(' СРЕДА')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)

    paragraph = doc3.add_paragraph('STYLE D ')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)

    for el_D in lst_D3:
        # перебираем программы
        for i, el_Pr in enumerate(el_D):

            # сохранение телепрограмм
            lst_tmp = save_prog(doc3, docN, el_Pr, el_D, i, str_prog3, str_progN)
            str_prog3 = lst_tmp[0]
            str_progN = lst_tmp[1]

            print(('Сохранение телепрограмм Среды - ' + progressSpin(progressInt)).ljust(60, " ") , end='\r')
            progressInt +=  1


    # четверг
    str_progN = str_progN + '-------------------------------\n' + ' ЧЕТВЕРГ\n'
    str_prog4 = 'STYLE D \n'
    # doc
    paragraph = docN.add_paragraph('-------------------------------')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)
    paragraph = docN.add_paragraph(' ЧЕТВЕРГ')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)

    paragraph = doc4.add_paragraph('STYLE D ')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)

    for el_D in lst_D4:
        # перебираем программы
        for i, el_Pr in enumerate(el_D):

            # сохранение телепрограмм
            lst_tmp = save_prog(doc4, docN, el_Pr, el_D, i, str_prog4, str_progN)
            str_prog4 = lst_tmp[0]
            str_progN = lst_tmp[1]

            print(('Сохранение телепрограмм Четверга - ' + progressSpin(progressInt)).ljust(60, " ") , end='\r')
            progressInt +=  1


    # пятница
    str_progN = str_progN + '-------------------------------\n' + ' ПЯТНИЦА\n'
    str_prog5 = 'STYLE D \n'
    # doc
    paragraph = docN.add_paragraph('-------------------------------')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)
    paragraph = docN.add_paragraph(' ПЯТНИЦА')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)

    paragraph = doc5.add_paragraph('STYLE D ')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)

    for el_D in lst_D5:
        # перебираем программы
        for i, el_Pr in enumerate(el_D):

            # сохранение телепрограмм
            lst_tmp = save_prog(doc5, docN, el_Pr, el_D, i, str_prog5, str_progN)
            str_prog5 = lst_tmp[0]
            str_progN = lst_tmp[1]

            print(('Сохранение телепрограмм Пятницы - ' + progressSpin(progressInt)).ljust(60, " ") , end='\r')
            progressInt +=  1


    # суббота
    str_progN = str_progN + '-------------------------------\n' + ' СУББОТА\n'
    str_prog6 = 'STYLE D \n'
    # doc
    paragraph = docN.add_paragraph('-------------------------------')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)
    paragraph = docN.add_paragraph(' СУББОТА')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)

    paragraph = doc6.add_paragraph('STYLE D ')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)

    for el_D in lst_D6:
        # перебираем программы
        for i, el_Pr in enumerate(el_D):

            # сохранение телепрограмм
            lst_tmp = save_prog(doc6, docN, el_Pr, el_D, i, str_prog6, str_progN)
            str_prog6 = lst_tmp[0]
            str_progN = lst_tmp[1]

            print(('Сохранение телепрограмм Субботы - ' + progressSpin(progressInt)).ljust(60, " ") , end='\r')
            progressInt +=  1


    # воскресенье
    str_progN = str_progN + '-------------------------------\n' + ' ВОСКРЕСЕНЬЕ\n'
    str_prog7 = 'STYLE D \n'
    # doc
    paragraph = docN.add_paragraph('-------------------------------')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)
    paragraph = docN.add_paragraph(' ВОСКРЕСЕНЬЕ')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)

    paragraph = doc7.add_paragraph('STYLE D ')
    paragraph.paragraph_format.space_before = Mm(0)
    paragraph.paragraph_format.space_after = Mm(0)
    for el_D in lst_D7:
        # перебираем программы
        for i, el_Pr in enumerate(el_D):

            # сохранение телепрограмм
            lst_tmp = save_prog(doc7, docN, el_Pr, el_D, i, str_prog7, str_progN)
            str_prog7 = lst_tmp[0]
            str_progN = lst_tmp[1]

            print(('Сохранение телепрограмм Воскресенья - ' + progressSpin(progressInt)).ljust(60, " ") , end='\r')
            progressInt +=  1

    # сохраняем данные 
    try:
        with open(path_prog_out +'REZULT.REZ', 'w') as file_w:
            file_w.writelines(str_progN)
    except:
        print('Файл REZULT.REZ заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt)).ljust(60, " ")  , end='\r')
    progressInt +=  1

    try:
        with open(path_prog_out +'1.REZ', 'w') as file_w:
            file_w.writelines(str_prog1)
    except:
        print('Файл 1.REZ заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt) ).ljust(60, " ") , end='\r')
    progressInt +=  1

    try:
        with open(path_prog_out +'2.REZ', 'w') as file_w:
            file_w.writelines(str_prog2)
    except:
        print('Файл 2.REZ заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt) ).ljust(60, " ") , end='\r')
    progressInt +=  1

    try:
        with open(path_prog_out +'3.REZ', 'w') as file_w:
            file_w.writelines(str_prog3)
    except:
        print('Файл 3.REZ заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt) ).ljust(60, " ")  , end='\r')
    progressInt +=  1

    try:
        with open(path_prog_out +'4.REZ', 'w') as file_w:
            file_w.writelines(str_prog4)
    except:
        print('Файл 4.REZ заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt)).ljust(60, " ")  , end='\r')
    progressInt +=  1

    try:
        with open(path_prog_out +'5.REZ', 'w') as file_w:
            file_w.writelines(str_prog5)
    except:
        print('Файл 5.REZ заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt)).ljust(60, " ")  , end='\r')
    progressInt +=  1

    try:
        with open(path_prog_out +'6.REZ', 'w') as file_w:
            file_w.writelines(str_prog6)
    except:
        print('Файл 6.REZ заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt) ).ljust(60, " ") , end='\r')
    progressInt +=  1

    try:
        with open(path_prog_out +'7.REZ', 'w') as file_w:
            file_w.writelines(str_prog7)
    except:
        print('Файл 7.REZ заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt) ).ljust(60, " ") , end='\r')
    progressInt +=  1

    #doc
    try:
        doc1.save(path_prog_out + '1.docx')
    except:
        print('Файл 1.docx заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt)).ljust(60, " ")  , end='\r')
    progressInt +=  1

    try:
        doc2.save(path_prog_out + '2.docx')
    except:
        print('Файл 2.docx заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt) ).ljust(60, " ") , end='\r')
    progressInt +=  1

    try:
        doc3.save(path_prog_out + '3.docx')
    except:
        print('Файл 3.docx заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt) ).ljust(60, " ") , end='\r')
    progressInt +=  1

    try:
        doc4.save(path_prog_out + '4.docx')
    except:
        print('Файл 1.docx заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt) ).ljust(60, " ") , end='\r')
    progressInt +=  1

    try:
        doc5.save(path_prog_out + '4.docx')
    except:
        print('Файл 1.docx заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt) ).ljust(60, " ") , end='\r')
    progressInt +=  1

    try:
        doc6.save(path_prog_out + '5.docx')
    except:
        print('Файл 5.docx заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt)).ljust(60, " ")  , end='\r')
    progressInt +=  1

    try:
        doc1.save(path_prog_out + '6.docx')
    except:
        print('Файл 7.docx заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt) ).ljust(60, " ") , end='\r')
    progressInt +=  1

    try:
        doc7.save(path_prog_out + '7.docx')
    except:
        print('Файл 1.docx заблокирован для вывода списка каналов!')

    print(('Сохранение результатов в файлы - ' + progressSpin(progressInt) ).ljust(60, " ") , end='\r')
    progressInt +=  1

    try:
        docN.save(path_prog_out + 'Rezult.docx')
    except:
        print('Файл Rezult.docx заблокирован для вывода списка каналов!')

    print('Сохранение результатов в файлы - ВЫПОЛНЕНО!')




# проверяем и делаем деКапсЛок
def deCapsLock(str_dcl):

    # ищем оставшиеся апострофы и меняем на елочки
    fl_search = True
    for pos_ap, sub_ap in enumerate(str_dcl):
        if sub_ap == '"':
            if pos_ap==0:
                str_dcl = '«' + str_dcl[1:]
            elif pos_ap>0 and pos_ap<len(str_dcl)-1:
                if str_dcl[pos_ap+1] == ' ':
                    str_dcl = str_dcl[:pos_ap] + '»' + str_dcl[pos_ap+1:]
                else:
                    str_dcl = str_dcl[:pos_ap] + '«' + str_dcl[pos_ap+1:]
            elif pos_ap == len(str_dcl)-1:
                str_dcl = str_dcl[:-1] + '»'




    str_dcl = str_dcl.strip()
    fl_caps = True
    fl_CapsWord = False
    lst_dcl = str_dcl.split(' ')
    tmp_str = ''
    for i, el in enumerate(lst_dcl):

        if lst_dcl[i]=='':
            continue

        # если первое слово и заглавные то устанавливаем первыю заглавную
        if i==0 and fl_caps:
            if lst_dcl[i][0].isalpha():
                tmp_str = lst_dcl[i][0].upper()  + lst_dcl[i][1:].lower()
            elif len(lst_dcl[i])>1:
                tmp_str = lst_dcl[i][:2].upper()  + lst_dcl[i][2:].lower()
            else:
                tmp_str = lst_dcl[i]
            continue

        #проверяем исключения по Капслоку
        for el_cpsl in lst_CapsWord:
            pos_cpsl = el.upper().find(el_cpsl.upper())
            if el.upper() == el_cpsl.upper():   # pos_cpsl>-1:
                if len(tmp_str.strip())>0:
                    tmp_str = tmp_str + ' ' + el_cpsl
                else:
                    tmp_str = el_cpsl
                fl_CapsWord = True
                break
            else:
                fl_CapsWord = False

        # если попалось исключение по CapsWord берем следующее слово
        if fl_CapsWord:
            continue

        fl_caps = True
        # по буквам проверяем слово на заглавные
        for j, sub_el in enumerate(el):
            if sub_el.isalpha() and fl_caps:
                if not sub_el.istitle():
                    # если попалась не заглавная сбрасываем флаг
                    fl_caps = False
                    break


        # если слово начинается с кавычек
        if i>0 and (lst_dcl[i][0]== '«' or lst_dcl[i][0]== '\'') and fl_caps:
            tmp_str = tmp_str + ' ' +  lst_dcl[i][0].upper()  + lst_dcl[i][1].upper()  + lst_dcl[i][2:].lower()

        # если предыдущее слово заканчивается точкой
        elif i>0 and (not lst_dcl[i-1]==''):
            if lst_dcl[i-1][-1]=='.':
                tmp_str = tmp_str + ' ' + lst_dcl[i][0].upper()  + lst_dcl[i][1:].lower()
            else:
                tmp_str = tmp_str + ' ' + lst_dcl[i].lower()

        # все прописные
        else:
            tmp_str = tmp_str + ' ' + lst_dcl[i].lower()

    # if fl_caps:
    #     str_dcl = str_dcl[0] + str_dcl[1:].lower()
    str_dcl = tmp_str.strip()

    return str_dcl




def del_dubl_prog():
    progressInt = 0

    # перебираем программы понедельника
    for k, el_prog in enumerate(lst_D1):
        for i, el in enumerate(el_prog):
            if i>0:
                # собираем дубликаты одинаковых строк
                lst_el = el.split('|',1)[1].split('^')[0].strip()
                lst_del = []
                for j, el_seek in enumerate(el_prog):
                    if j>i:
                        lst_el_seek = el_seek.split('|', 1)[1].split('^')[0].strip()
                        if lst_el.strip()==lst_el_seek.strip():
                            # найден дубль в j для i
                            lst_D1[k][i] = lst_D1[k][i].split("|")[0] + ', ' + lst_D1[k][j].split('|')[0] + '|' + lst_D1[k][i].split("|")[1]
                            lst_del.append(j)


                # # собираем дубликаты названий в ковычках
                # # сравнение внутри кавычек до первой точки
                # lst_name_el = str_kav_dot(lst_el)
                # if not (lst_name_el == ''):
                #     for jn, el_seekn in enumerate(el_prog):
                #         if jn>1:
                #             lst_eln_seek = el_seekn.split('|', 1)[1]
                #             if lst_eln_seek.find(lst_name_el)>-1:
                #                 # найден дубль в jn для i
                #                 lst_D1[k][i] = lst_D1[k][i].split("|")[0] + ', ' + lst_D1[k][j].split('|')[0] + '|' + lst_D1[k][i].split("|")[1]
                #                 lst_del.append(jn)

                # удаляем собранные дубликаты в обратном порядке
                for n in reversed(lst_del):
                    del lst_D1[k][n]







    print('Собираем дубликаты программ - ' + progressSpin(progressInt) , end='\r')
    progressInt +=  1

    # перебираем программы Вторника
    for k, el_prog in enumerate(lst_D2):
        for i, el in enumerate(el_prog):
            if i>0:
                lst_el = el.split('|',1)[1]
                lst_del = []
                for j, el_seek in enumerate(el_prog):
                    if j>i:
                        lst_el_seek = el_seek.split('|',1)[1]
                        if lst_el==lst_el_seek:
                            # найден дубль в j для i
                            lst_D2[k][i] = lst_D2[k][i].split("|")[0] + ', ' + lst_D2[k][j].split('|')[0] + '|' + lst_D2[k][i].split("|")[1]
                            lst_del.append(j)

                for n in reversed(lst_del):
                    del lst_D2[k][n]

    print('Собираем дубликаты программ - ' + progressSpin(progressInt) , end='\r')
    progressInt +=  1

    # перебираем программы среда
    for k, el_prog in enumerate(lst_D3):
        for i, el in enumerate(el_prog):
            if i>0:
                lst_el = el.split('|',1)[1]
                lst_del = []
                for j, el_seek in enumerate(el_prog):
                    if j>i:
                        lst_el_seek = el_seek.split('|',1)[1]
                        if lst_el==lst_el_seek:
                            # найден дубль в j для i
                            lst_D3[k][i] = lst_D3[k][i].split("|")[0] + ', ' + lst_D3[k][j].split('|')[0] + '|' + lst_D3[k][i].split("|")[1]
                            lst_del.append(j)

                for n in reversed(lst_del):
                    del lst_D3[k][n]

    print('Собираем дубликаты программ - ' + progressSpin(progressInt) , end='\r')
    progressInt +=  1

     # перебираем программы четверг
    for k, el_prog in enumerate(lst_D4):
        for i, el in enumerate(el_prog):
            if i>0:
                lst_el = el.split('|',1)[1]
                lst_del = []
                for j, el_seek in enumerate(el_prog):
                    if j>i:
                        lst_el_seek = el_seek.split('|',1)[1]
                        if lst_el==lst_el_seek:
                            # найден дубль в j для i
                            lst_D4[k][i] = lst_D4[k][i].split("|")[0] + ', ' + lst_D4[k][j].split('|')[0] + '|' + lst_D4[k][i].split("|")[1]
                            lst_del.append(j)

                for n in reversed(lst_del):
                    del lst_D4[k][n]

    print('Собираем дубликаты программ - ' + progressSpin(progressInt) , end='\r')
    progressInt +=  1

    # перебираем программы пятница
    for k, el_prog in enumerate(lst_D5):
        for i, el in enumerate(el_prog):
            if i>0:
                lst_el = el.split('|',1)[1]
                lst_del = []
                for j, el_seek in enumerate(el_prog):
                    if j>i:
                        lst_el_seek = el_seek.split('|',1)[1]
                        if lst_el==lst_el_seek:
                            # найден дубль в j для i
                            lst_D5[k][i] = lst_D5[k][i].split("|")[0] + ', ' + lst_D5[k][j].split('|')[0] + '|' + lst_D5[k][i].split("|")[1]
                            lst_del.append(j)

                for n in reversed(lst_del):
                    del lst_D5[k][n]

    print('Собираем дубликаты программ - ' + progressSpin(progressInt) , end='\r')
    progressInt +=  1

    # перебираем программы суббота
    for k, el_prog in enumerate(lst_D6):
        for i, el in enumerate(el_prog):
            if i>0:
                lst_el = el.split('|',1)[1]
                lst_del = []
                for j, el_seek in enumerate(el_prog):
                    if j>i:
                        lst_el_seek = el_seek.split('|',1)[1]
                        if lst_el==lst_el_seek:
                            # найден дубль в j для i
                            lst_D6[k][i] = lst_D6[k][i].split("|")[0] + ', ' + lst_D6[k][j].split('|')[0] + '|' + lst_D6[k][i].split("|")[1]
                            lst_del.append(j)

                for n in reversed(lst_del):
                    del lst_D6[k][n]

    print('Собираем дубликаты программ - ' + progressSpin(progressInt) , end='\r')
    progressInt +=  1

    # перебираем программы воскресенье
    for k, el_prog in enumerate(lst_D7):
        for i, el in enumerate(el_prog):
            if i>0:
                lst_el = el.split('|',1)[1]
                lst_del = []
                for j, el_seek in enumerate(el_prog):
                    if j>i:
                        lst_el_seek = el_seek.split('|',1)[1]
                        if lst_el==lst_el_seek:
                            # найден дубль в j для i
                            lst_D7[k][i] = lst_D7[k][i].split("|")[0] + ', ' + lst_D7[k][j].split('|')[0] + '|' + lst_D7[k][i].split("|")[1]
                            lst_del.append(j)

                for n in reversed(lst_del):
                    del lst_D7[k][n]

    print('Собираем дубликаты программ - ВЫПОНЕНО!' )
    progressInt +=  1


# делаем анализ строки программы 
# проводим замены и синтезируем строку 
# str_sub_repl - строка замены 
# str_sub_vozrast - строка возрастное ограничение
# str_sub_name_prog - строка программы внутри кавычек елочкой

def analiz_prog_day_in(str_prog,                # анализируемая строка
                       name_Channel,            # название канала
                       fl_Caps):                # обработка Капса 

    # определяем переменные для синтеза программы
    str_sub_vozrast = ''         # возрастная категория, при синтезе ставится в конец строки
    str_sub_repl = ''            # вид передачи, при синтезе ставится в начало строки
    str_sub_name_prog = ''       # название передачи в кавычках


    fl_stop_word = False         # флаг найденного стоп слова (строка не обрабатывается)
    fl_stop_AP = False           # флаг найденной авторской программы (строка обрабатывается по справочнику)
    fl_stop_Repl = False         # флаг найденной замены

    #ищем и удаляем предварительно киностудии ленфильм мосфильм
    nameFilm = '"ЛЕНФИЛЬМ"'
    posFilm = str_prog.upper().find(nameFilm)
    if posFilm > -1:
        str_prog = str_prog[:posFilm] + ' ' + str_prog[posFilm + len(nameFilm) :]

    nameFilm = '"МОСФИЛЬМ"'
    posFilm = str_prog.upper().find(nameFilm)
    if posFilm > -1:
        str_prog = str_prog[:posFilm] + ' ' + str_prog[posFilm + len(nameFilm) :]

    find_v = False
    find_v1 = False
    find_v2 = False
    # сканируем в поиске возрастной категории
    for j, el_v in enumerate(vozrast):
        if el_v in str_prog:
            str_sub_vozrast = vozrast[j]
            find_v = True
            break
    if not find_v:
        for j, el_v in enumerate(vozrast1):
            if el_v in str_prog:
                str_sub_vozrast = vozrast1[j]
                find_v1 = True
                break
    if not find_v1 and not find_v:
        for j, el_v in enumerate(vozrast2):
            if el_v in str_prog:
                str_sub_vozrast = vozrast2[j]
                find_v2 = True
                break

            
    # вырезаем из строки 
    # и сохраняем возрастной индекс str_sub_vozrast
    if not str_sub_vozrast=='':
        if find_v2:
            str_prog = str_prog.replace(str_sub_vozrast, '').strip()
            str_sub_vozrast = str_sub_vozrast.replace(' ', '').strip()

        if find_v1 and not find_v2:
            str_prog = str_prog.replace(str_sub_vozrast, '').strip()
            str_sub_vozrast = str_sub_vozrast.replace('[', ' ').replace(']', '').strip()

        if find_v and not find_v1:
            str_prog = str_prog.replace(str_sub_vozrast, '').strip()
            str_sub_vozrast = str_sub_vozrast.replace('(', '').replace(')', '').strip()


        


    # определяем наличие стоп слова в строке
    for el in lst_StopWord:
        if str_prog.upper().find(el.upper()) > -1 :
            str_sub_repl = 'DEL'   # авторская программа (при синтезе удаляем)
            str_sub_name_prog = str_prog
            fl_stop_word = True
            break


    # определяем наличие авторской программы
    if not fl_stop_word:
        for el in lst_Programma:
            if str_prog.upper().find(el[0].upper()) > -1 and name_Channel.upper()==el[2].upper():
                str_sub_repl = 'DEL'   # авторская программа (при синтезе удаляем)
                str_sub_name_prog = el[1]
                fl_stop_AP = True
                break



    # определяем замену
    if (not fl_stop_word)  :
        # и сохраняем в переменной str_sub_repl
        for el in lst_Repl:
            pos_repl = str_prog.upper().find(el[0].upper())
            if str_prog.upper().find(el[0].upper()) > -1 and name_Channel.upper()==el[2].upper():
                if str_sub_name_prog == '':
                    str_sub_name_prog = 'DEL'   # при синтезе удаляем
                str_sub_repl = el[1]
                str_prog = str_prog[:pos_repl].strip() + ' ' + str_prog[pos_repl + len(el[0]) :].strip()
                fl_stop_Repl = True
                break
            else:
                str_sub_repl = ''



    # Название программы в кавычках
    # и вырезаем строку внутри кавычек ёлочек
    if (not fl_stop_word) and (not fl_stop_AP) :
        pos1 = str_prog.find('«')
        pos2 = str_prog.rfind('»')

        if not ((pos1 > -1) and (pos2 > pos1)):
            # если не найдены кавычки ёлочки ищем двойные апострофы
            pos1 = str_prog.find('"')
            pos2 = str_prog.rfind('"')

            if ((pos1 > -1) and (pos2 > pos1)):
                if str_sub_repl == '':
                    str_sub_repl = 'DEL'   # при синтезе удаляем
                if fl_Caps:
                    str_sub_name_prog = '«' + deCapsLock(str_prog[pos1+1:pos2]) + '»'
                else:
                    str_sub_name_prog = '«' + str_prog[pos1+1:pos2] + '»'
                str_prog = str_prog[:pos1] + ' ' + str_prog[pos2+1:]
            else:
                str_sub_name_prog = ''

        else:
            if ((pos1 > -1) and (pos2 > pos1)):
                if str_sub_repl == '':
                    str_sub_repl = 'DEL'   # при синтезе удаляем
                if fl_Caps:
                    str_sub_name_prog = '«' + deCapsLock(str_prog[pos1+1:pos2]) + '»'
                else:
                    str_sub_name_prog = '«' + str_prog[pos1+1:pos2] + '»'
                str_prog = str_prog[:pos1] + ' ' + str_prog[pos2+1:]
            else:
                str_sub_name_prog = ''




    if not fl_stop_word:
        # синтезируем строку программы
        fl_sintez = 0
        str_sintez = ''
        if not str_sub_repl == '':
            if not str_sub_repl == 'DEL':
                str_sintez =  str_sub_repl
            fl_sintez += 1


        if not str_sub_name_prog == '':
            if not str_sub_name_prog == 'DEL':
                str_sintez = str_sintez + ' ' + str_sub_name_prog
            fl_sintez += 1

        if str_sintez == '':
            if fl_Caps:
                str_sintez =  deCapsLock(str_prog)
            else:
                str_sintez =  str_prog
        else:
            if fl_sintez < 2:
                if fl_stop_AP:
                    str_sintez = str_sintez
                else:
                    if name_Channel.upper() == 'СПАС':  # в СПАСе названия фильмов без кавычек
                        if fl_Caps:
                            str_sintez = str_sintez + ' «' + deCapsLock(str_prog) + '»'
                        else:
                            str_sintez = str_sintez + ' «' + str_prog + '»'
                    else:
                        if fl_Caps:
                            str_sintez = str_sintez + ' ' + deCapsLock(str_prog)
                        else:
                            str_sintez = str_sintez + ' ' + str_prog
    else:
        if fl_Caps:
            str_sintez = deCapsLock(str_prog)
        else:
            str_sintez = str_prog


    # удаляем по справочнику
    if (not fl_stop_word) and (not fl_stop_AP) :
        # ищем и удаляем слова по списку удаления
        str_sub_remove = ''
        for i, el in enumerate(lst_Remove):
            if el in str_sintez:
                pos_Rem = str_sintez.upper().find(el.upper())
                str_sintez = str_sintez[:pos_Rem] + ' ' + str_sintez[pos_Rem + len(el):]


    str_sintez = str_sintez.replace(' .', '.').replace(' .', '.').strip()

    # добавляем возрастное ограничение
    if not str_sub_vozrast == '':
        str_sintez = str_sintez + ' ^' + str_sub_vozrast

    Rezult = str_sintez





    return Rezult


# def remove_trash():
#     print()



# основное тело программы
def main():

    print('Парсер для обработки программ телеканалов ')
    print('Макеев Петр тел. 8-912-34-161-34 ')
    print('-----------------------------------------------')


    # проверяем дату
    # if datetime.date.today().month>10 and datetime.date.today().day>20:
    #     input('Закончился демо-режим программы, нажмите Enter для завершения.')
    #     exit()

    # определяем окружение
    path_prog = os.getcwd()
    if platform.system() == 'Windows':
        path_prog_in = path_prog + '\\in\\'
        path_prog_out = path_prog + '\\out\\'
        path_prog = path_prog + '\\'
    elif platform.system() == 'Linux':
        path_prog_in = path_prog + '/in/'
        path_prog_out = path_prog + '/out/'
        path_prog = path_prog + '/'
    else:
        input('Не определена операционная система! Нажмите любую клавишу.')
        exit()


    # сделать запрос поправки на часовой пояс, по умолчанию поставить 2 часа
    #read_divH()

   # считываем тхт файлы
    txt_to_list(path_prog, path_prog_in)

    # создаем заготовки списков программ по дням
    fill_Day()

    # считываем каналы 
    txt_to_prog(path_prog_in)

    # analiz_prog()
    analiz_prog()

    # сводим дубликаты программ в одну строку
    del_dubl_prog()

     # сохраняем программы в файлы
    exp_prog(path_prog_out)

    input('Телепрограммы обработаны, результаты в папке OUT, нажмите Enter для завершения.')
    # print('w')   


# точка входа.
if __name__ == '__main__':
    main()


